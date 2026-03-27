"""
Resume tailoring via Claude API.

Optimisations:
- Prompt caching: system prompt + base resume are cached across all calls in a run
  (saves ~90% on those tokens after the first call)
- Batch API mode: when USE_BATCH_API=true, all tailoring requests are submitted as
  a single batch at 50% discount; results are processed later via `job-agent batch-process`
"""

import subprocess
from pathlib import Path

import anthropic
from docx import Document

from job_agent.config import settings
from job_agent.models import Job

SOFFICE = "/Applications/LibreOffice.app/Contents/MacOS/soffice"

TAILORABLE_STYLES = {"Body A", "Body B"}
STOP_HEADINGS = {"education", "skills", "skilss"}

SYSTEM_PROMPT = """You are a professional resume writer helping a job seeker tailor their resume.

STRICT RULES — violations will disqualify the output:
1. You MAY reorder bullet points within each job to surface the most relevant experience first.
2. You MAY adjust wording of existing bullets to incorporate keywords from the job description.
3. You MUST NOT add any experience, skills, projects, or accomplishments not present in the original.
4. You MUST NOT change job titles, company names, dates, or education details.
5. You MUST return exactly the same number of paragraphs as provided, in the same index order.

Input format:
  <index>: <paragraph text>

Output format — return ONLY lines in this format, one per paragraph:
  <index>: <tailored paragraph text>

After all paragraphs, add a line with exactly "---CHANGES---" followed by a brief summary of changes."""

_client: anthropic.Anthropic | None = None


def _get_client() -> anthropic.Anthropic:
    global _client
    if _client is None:
        _client = anthropic.Anthropic(api_key=settings.anthropic_api_key)
    return _client


# ── docx helpers ─────────────────────────────────────────────────────────────

def _get_tailorable_paragraphs(doc: Document) -> list[tuple[int, str]]:
    """Returns (doc_index, text) for body paragraphs in Profile and Experience sections only."""
    result = []
    stop = False
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading" and p.text.strip().lower() in STOP_HEADINGS:
            stop = True
        if stop:
            continue
        if p.style.name in TAILORABLE_STYLES and p.text.strip():
            result.append((i, p.text))
    return result


def _replace_paragraph_text(para, new_text: str) -> None:
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def _write_tailored_docx(base_docx_path: Path, replacements: dict[int, str], output_path: Path) -> None:
    doc = Document(str(base_docx_path))
    for idx, new_text in replacements.items():
        _replace_paragraph_text(doc.paragraphs[idx], new_text)
    doc.save(str(output_path))


def _convert_to_pdf(docx_path: Path) -> Path:
    subprocess.run(
        [SOFFICE, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
        check=True,
        capture_output=True,
    )
    return docx_path.with_suffix(".pdf")


def _parse_response(full_response: str) -> tuple[dict[int, str], str]:
    """Parse Claude's indexed response into a replacements dict and change notes."""
    if "---CHANGES---" in full_response:
        content, notes = full_response.split("---CHANGES---", 1)
    else:
        content, notes = full_response, "No changes documented."

    replacements: dict[int, str] = {}
    for line in content.strip().splitlines():
        if ": " in line:
            raw_idx, _, text = line.partition(": ")
            try:
                replacements[int(raw_idx.strip())] = text.strip()
            except ValueError:
                pass
    return replacements, notes.strip()


def _build_messages(job: Job, input_block: str, base_resume_text: str) -> list[dict]:
    """Build the messages array, keeping the base resume in a cached block."""
    return [
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": f"Base Resume (for context — do not modify structure):\n{base_resume_text}",
                    "cache_control": {"type": "ephemeral"},
                },
                {
                    "type": "text",
                    "text": (
                        f"Job Title: {job.title}\n"
                        f"Company: {job.company}\n\n"
                        f"Job Description:\n{job.description}\n\n"
                        f"Resume paragraphs to tailor:\n{input_block}"
                    ),
                },
            ],
        }
    ]


# ── public API ────────────────────────────────────────────────────────────────

def tailor_resume(job: Job, base_docx_path: Path) -> tuple[Path, Path, str]:
    """
    Synchronous tailoring with prompt caching.
    Returns (tailored_docx_path, tailored_pdf_path, notes_on_changes).
    """
    doc = Document(str(base_docx_path))
    tailorable = _get_tailorable_paragraphs(doc)
    base_resume_text = "\n".join(p.text for _, p_idx in [(None, i) for i in range(len(doc.paragraphs))]
                                 for p in [doc.paragraphs[p_idx]])
    input_block = "\n".join(f"{idx}: {text}" for idx, text in tailorable)

    # Re-extract base resume as plain text for the cached context block
    base_text_lines = [p.text for p in doc.paragraphs]
    base_resume_text = "\n".join(base_text_lines)

    response = _get_client().messages.create(
        model=settings.claude_model,
        max_tokens=4096,
        system=[
            {
                "type": "text",
                "text": SYSTEM_PROMPT,
                "cache_control": {"type": "ephemeral"},
            }
        ],
        messages=_build_messages(job, input_block, base_resume_text),
    )

    replacements, notes = _parse_response(response.content[0].text)

    safe_company = job.company.replace(" ", "_").replace("/", "-")
    stem = f"{job.id}_{safe_company}"
    docx_path = settings.tailored_resume_dir / f"{stem}.docx"
    _write_tailored_docx(base_docx_path, replacements, docx_path)
    pdf_path = _convert_to_pdf(docx_path)

    return docx_path, pdf_path, notes


def build_batch_request(job: Job, base_docx_path: Path) -> dict:
    """Build a single batch request dict for the Messages Batch API."""
    doc = Document(str(base_docx_path))
    tailorable = _get_tailorable_paragraphs(doc)
    base_resume_text = "\n".join(p.text for p in doc.paragraphs)
    input_block = "\n".join(f"{idx}: {text}" for idx, text in tailorable)

    return {
        "custom_id": str(job.id),
        "params": {
            "model": settings.claude_model,
            "max_tokens": 4096,
            "system": [
                {
                    "type": "text",
                    "text": SYSTEM_PROMPT,
                    "cache_control": {"type": "ephemeral"},
                }
            ],
            "messages": _build_messages(job, input_block, base_resume_text),
        },
    }


def process_batch_result(job: Job, response_text: str, base_docx_path: Path) -> tuple[Path, Path, str]:
    """
    Convert a batch API response text into docx + pdf files.
    Same output as tailor_resume().
    """
    replacements, notes = _parse_response(response_text)

    safe_company = job.company.replace(" ", "_").replace("/", "-")
    stem = f"{job.id}_{safe_company}"
    docx_path = settings.tailored_resume_dir / f"{stem}.docx"
    _write_tailored_docx(base_docx_path, replacements, docx_path)
    pdf_path = _convert_to_pdf(docx_path)

    return docx_path, pdf_path, notes